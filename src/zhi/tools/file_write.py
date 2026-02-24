"""File write tool for zhi."""

from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from typing import Any, ClassVar

import docx
import openpyxl

from zhi.tools.base import BaseTool

_MAX_OUTPUT_SIZE = 50 * 1024  # 50KB tool output cap


class FileWriteTool(BaseTool):
    """Write a new file to the output directory.

    Supported formats by extension:
    - .md, .txt: plain text string
    - .json: JSON pass-through
    - .csv: {"headers": [...], "rows": [[...]]}
    - .xlsx: {"sheets": [{"name": "...", "headers": [...], "rows": [[...]]}]}
    - .docx: {"content": "markdown string"}
    """

    name: ClassVar[str] = "file_write"
    description: ClassVar[str] = (
        "Write a new file to the output directory (zhi-output/). "
        "Supports: .md, .txt (plain text), .json (any JSON), "
        ".csv ({headers, rows}), "
        ".xlsx ({sheets: [{name, headers, rows}]}), "
        '.docx ({content: "markdown"}). '
        "Cannot overwrite existing files."
    )
    parameters: ClassVar[dict[str, Any]] = {
        "type": "object",
        "properties": {
            "path": {
                "type": "string",
                "description": (
                    "Relative path within the output directory for the new file."
                ),
            },
            "content": {
                "description": (
                    "File content. For .md/.txt: plain text string. "
                    "For .json: any JSON value. "
                    'For .csv: {"headers": [...], "rows": [[...]]}. '
                    "For .xlsx: {sheets: [{name, headers, rows}]}. "
                    "For .docx: {content: markdown string}."
                ),
            },
        },
        "required": ["path", "content"],
    }
    risky: ClassVar[bool] = True

    def __init__(self, output_dir: Path | None = None) -> None:
        self._output_dir = output_dir or Path.cwd() / "zhi-output"

    def execute(self, **kwargs: Any) -> str:
        path_str: str = kwargs.get("path", "")
        content: Any = kwargs.get("content")

        if not path_str:
            return "Error: 'path' parameter is required."
        if content is None:
            return "Error: 'content' parameter is required."

        file_path = Path(path_str)

        # Reject absolute paths
        if file_path.is_absolute():
            return "Error: Absolute paths are not allowed. Use a relative path."

        # Reject path traversal
        if ".." in file_path.parts:
            return "Error: Path traversal ('..') is not allowed."

        # Ensure output directory exists
        self._output_dir.mkdir(parents=True, exist_ok=True)

        target = self._output_dir / file_path

        # Resolve symlinks and verify still within output dir
        try:
            resolved = target.resolve()
            output_resolved = self._output_dir.resolve()
            if resolved != output_resolved and not resolved.is_relative_to(
                output_resolved
            ):
                return (
                    "Error: Resolved path is outside the "
                    "output directory. Possible symlink attack."
                )
        except (OSError, ValueError):
            return "Error: Invalid file path."

        # No-overwrite check
        if resolved.exists():
            return (
                f"Error: File already exists: {path_str}. "
                "Cannot overwrite existing files."
            )

        # Create parent directories
        resolved.parent.mkdir(parents=True, exist_ok=True)

        ext = file_path.suffix.lower()

        try:
            if ext in (".md", ".txt"):
                return self._write_text(resolved, content, path_str)
            elif ext == ".json":
                return self._write_json(resolved, content, path_str)
            elif ext == ".csv":
                return self._write_csv(resolved, content, path_str)
            elif ext == ".xlsx":
                return self._write_xlsx(resolved, content, path_str)
            elif ext == ".docx":
                return self._write_docx(resolved, content, path_str)
            else:
                # Default to text
                return self._write_text(resolved, content, path_str)
        except OSError as exc:
            return f"Error: Could not write file: {exc}"

    def _write_text(self, path: Path, content: Any, rel_path: str) -> str:
        text = str(content)
        path.write_text(text, encoding="utf-8")
        size = path.stat().st_size
        return self._success(rel_path, size)

    def _write_json(self, path: Path, content: Any, rel_path: str) -> str:
        if isinstance(content, str):
            # Validate it's valid JSON
            try:
                parsed = json.loads(content)
                text = json.dumps(parsed, indent=2, ensure_ascii=False)
            except json.JSONDecodeError:
                text = content
        else:
            text = json.dumps(content, indent=2, ensure_ascii=False)
        path.write_text(text, encoding="utf-8")
        size = path.stat().st_size
        return self._success(rel_path, size)

    def _write_csv(self, path: Path, content: Any, rel_path: str) -> str:
        if isinstance(content, str):
            try:
                content = json.loads(content)
            except json.JSONDecodeError:
                # Treat as raw CSV text
                path.write_text(content, encoding="utf-8")
                size = path.stat().st_size
                return self._success(rel_path, size)

        if not isinstance(content, dict):
            return 'Error: CSV content must be {"headers": [...], "rows": [[...]]}.'

        headers = content.get("headers", [])
        rows = content.get("rows", [])

        buf = io.StringIO()
        writer = csv.writer(buf)
        if headers:
            writer.writerow(headers)
        writer.writerows(rows)

        path.write_text(buf.getvalue(), encoding="utf-8")
        size = path.stat().st_size
        return self._success(rel_path, size)

    def _write_xlsx(self, path: Path, content: Any, rel_path: str) -> str:
        if isinstance(content, str):
            try:
                content = json.loads(content)
            except json.JSONDecodeError:
                return "Error: XLSX content must be valid JSON with sheets data."

        if not isinstance(content, dict):
            return (
                "Error: XLSX content must be "
                '{"sheets": [{"name": ..., "headers": '
                '[...], "rows": [[...]]}]}.'
            )

        wb = openpyxl.Workbook()
        sheets = content.get("sheets", [])

        for i, sheet_data in enumerate(sheets):
            if i == 0:
                ws = wb.active
                ws.title = sheet_data.get("name", "Sheet1")
            else:
                title = sheet_data.get("name", f"Sheet{i + 1}")
                ws = wb.create_sheet(title=title)

            headers = sheet_data.get("headers", [])
            rows = sheet_data.get("rows", [])

            if headers:
                ws.append(headers)
            for row in rows:
                ws.append(row)

        wb.save(str(path))
        size = path.stat().st_size
        return self._success(rel_path, size)

    def _write_docx(self, path: Path, content: Any, rel_path: str) -> str:
        if isinstance(content, str):
            try:
                content = json.loads(content)
            except json.JSONDecodeError:
                # Treat string as markdown content directly
                content = {"content": content}

        if not isinstance(content, dict):
            return 'Error: DOCX content must be {"content": "markdown string"}.'

        md_text = content.get("content", "")

        doc = docx.Document()
        # Simple paragraph splitting
        for line in md_text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped:
                doc.add_paragraph(stripped)

        doc.save(str(path))
        size = path.stat().st_size
        return self._success(rel_path, size)

    def _success(self, rel_path: str, size: int) -> str:
        size_str = f"{size / 1024:.1f}KB" if size >= 1024 else f"{size}B"
        result = f"File written: {rel_path} ({size_str})"
        if len(result) > _MAX_OUTPUT_SIZE:
            result = result[:_MAX_OUTPUT_SIZE] + "\n[output truncated at 50KB]"
        return result
